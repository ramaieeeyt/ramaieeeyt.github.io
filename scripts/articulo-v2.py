# -*- coding: utf-8 -*-
"""CONTEST-ARTICLE-V2.md -> .docx + .txt, solo los bloques en ingles."""
import io, re, os
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

BASE = '/Users/dez/Desktop/IEEE YT webpage'
src  = io.open(os.path.join(BASE, 'CONTEST-ARTICLE-V2.md'), encoding='utf-8').read()

ORDEN = ['TÍTULO', 'BLOQUE DE AUTORES', 'ABSTRACT', 'KEYWORDS',
         'I. INTRODUCTION', 'II. METHODOLOGY', 'III. IMPLEMENTATION',
         'IV. RESULTS', 'V. CLOSING REMARKS', 'VI. CONCLUSIONS', 'REFERENCES']

# --- extraer cada bloque ``` bajo su ## titulo ---
bloques = {}
for m in re.finditer(r'^## (.+?)\n(.*?)(?=^## |\Z)', src, re.S | re.M):
    titulo, resto = m.group(1).strip(), m.group(2)
    fence = re.search(r'```\n(.*?)```', resto, re.S)
    if not fence:
        continue
    cuerpo = fence.group(1)
    if titulo in ('TÍTULO', 'BLOQUE DE AUTORES'):
        # cada linea es su propio parrafo
        parrafos = [l.strip() for l in cuerpo.split('\n') if l.strip()]
    else:
        parrafos = [' '.join(x.split()) for x in cuerpo.split('\n\n') if x.strip()]
    bloques[titulo] = parrafos

faltan = [k for k in ORDEN if k not in bloques]
assert not faltan, 'faltan bloques: %s' % faltan

# --- la Tabla 1, anclada a su cabecera para no confundirla con las otras ---
tabla = []
for linea in src[src.index('| Criteria |'):].split('\n'):
    if not linea.startswith('|'):
        break
    if '---' in linea:
        continue
    tabla.append([c.strip() for c in linea.strip().strip('|').split('|')])
assert len(tabla) == 6 and len(tabla[0]) == 3, 'Tabla 1: %d filas' % len(tabla)

# ============================================================ DOCX
d = Document()
sec = d.sections[0]
sec.page_width, sec.page_height = Inches(8.5), Inches(11)          # Letter
for lado in ('top', 'bottom', 'left', 'right'):
    setattr(sec, '%s_margin' % lado, Inches(0.75))

n = d.styles['Normal']
n.font.name, n.font.size = 'Times New Roman', Pt(10)
pf = n.paragraph_format
pf.space_after, pf.line_spacing = Pt(7), 1.15

def p(texto, size=10, bold=False, italic=False, align=None, space_before=0):
    par = d.add_paragraph()
    par.alignment = align if align is not None else WD_ALIGN_PARAGRAPH.JUSTIFY
    par.paragraph_format.space_before = Pt(space_before)
    r = par.add_run(texto)
    r.font.size, r.bold, r.italic = Pt(size), bold, italic
    return par

C = WD_ALIGN_PARAGRAPH.CENTER
L = WD_ALIGN_PARAGRAPH.LEFT

titulo = bloques['TÍTULO']
p(titulo[0], size=20, bold=True, align=C)
p(titulo[1], size=14, italic=True, align=C)
for linea in bloques['BLOQUE DE AUTORES']:
    p(linea, size=10, align=C)

p('Abstract', size=12, bold=True, align=L, space_before=10)
for x in bloques['ABSTRACT']:
    p(x, italic=True)
p('Keywords', size=12, bold=True, align=L, space_before=6)
for x in bloques['KEYWORDS']:
    p(x, italic=True)

for clave in ORDEN[4:]:
    p(clave if clave != 'REFERENCES' else 'References',
      size=12, bold=True, align=L, space_before=10)
    for x in bloques[clave]:
        p(x, align=L if clave == 'REFERENCES' else None)

    if clave == 'IV. RESULTS':
        p('Table 1', size=10, bold=True, align=C, space_before=8)
        t = d.add_table(rows=0, cols=3)
        t.style, t.alignment = 'Table Grid', WD_TABLE_ALIGNMENT.CENTER
        for i, fila in enumerate(tabla):
            celdas = t.add_row().cells
            for j, txt in enumerate(fila):
                cp = celdas[j].paragraphs[0]
                cp.alignment = C if j == 1 else L
                cr = cp.add_run(txt)
                cr.font.size, cr.font.name = Pt(9), 'Times New Roman'
                cr.bold = (i == 0)
        p("Website evaluation criteria and the site's response to each.",
          size=9, italic=True, align=C)

d.save(os.path.join(BASE, 'CONTEST-ARTICLE-V2.docx'))

# ============================================================ TXT
out = []
out += [titulo[0], titulo[1], '']
out += bloques['BLOQUE DE AUTORES'] + ['']
for clave in ['ABSTRACT', 'KEYWORDS'] + ORDEN[4:]:
    out += ['=' * 70, clave, '=' * 70, '']
    for x in bloques[clave]:
        out += [x, '']
    if clave == 'IV. RESULTS':
        out += ['Table 1', '']
        for fila in tabla:
            out += [' | '.join(fila)]
        out += ['', "Website evaluation criteria and the site's response to each.", '']
io.open(os.path.join(BASE, 'CONTEST-ARTICLE-V2.txt'), 'w', encoding='utf-8')\
  .write('\n'.join(out))

palabras = sum(len(x.split()) for k in ORDEN for x in bloques[k])
print('bloques  :', len(bloques))
print('palabras :', palabras)
print('tabla    :', len(tabla), 'filas')
